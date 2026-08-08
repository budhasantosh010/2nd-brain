"""DOCX parser with paragraph/heading locators."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from second_brain.models import ParsedSegment
from second_brain.parsers.base import BaseParser


class DOCXParser(BaseParser):
    extensions = frozenset({".docx"})

    def parse(self, path: Path, source_id: str):  # type: ignore[no-untyped-def]
        document = Document(str(path))
        segments: list[ParsedSegment] = []
        heading = "document"
        buffer: list[str] = []
        first_para = 1

        def flush(last_para: int) -> None:
            nonlocal buffer, first_para
            body = "\n".join(buffer).strip()
            if body:
                segments.append(
                    ParsedSegment(
                        segment_id=f"{source_id}:seg:{len(segments)}",
                        text=body,
                        locator=f"{heading}; paragraphs {first_para}-{last_para}",
                        position=len(segments),
                    )
                )
            buffer = []

        for number, paragraph in enumerate(document.paragraphs, start=1):
            style = paragraph.style.name if paragraph.style else ""
            if style.lower().startswith("heading") and paragraph.text.strip():
                flush(number - 1)
                heading = f"heading: {paragraph.text.strip()}"
                first_para = number
            if paragraph.text.strip():
                buffer.append(paragraph.text)
        flush(len(document.paragraphs))
        if not segments:
            segments.append(
                ParsedSegment(segment_id=f"{source_id}:seg:0", text="", locator="document", position=0)
            )
        return self.document(
            path,
            source_id,
            segments,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
